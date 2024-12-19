import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

import os

doc = docx.Document()
a = doc.add_paragraph('                             תאריך 18/04/20 צבן מדיקל בע"מ RD41006441C קובץ: FINAL TEST REPORT - OXYGEN CONCENTRATOR'
                      'דו"ח סופיות(REV:05)')

run = doc.add_paragraph().add_run()
font = run.font

a.style = doc.styles.add_style('Style Name', WD_STYLE_TYPE.PARAGRAPH)
font = a.style.font
font.bold = True
font.size = Pt(9)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.4)

records = [
    ['                        PK444A05M', '                Oxygen Concentrator 5L', '                          OXYTEC 5S'],
    ["                        מס' סידורי מדחס", "                             ברקוד צבאן", "                       מס' סידורי למערכת"],
    ["SN:", "", "SN:"],
    ["", "", ""],
    ["                        מס' סידורי כרטיס", "", "                      מס' סידורי מכלי סינון"],
    ["SN:", "", "SN:"]
]
menuTable = doc.add_table(rows=1, cols=3)
menuTable.alignment = WD_TABLE_ALIGNMENT.CENTER
menuTable.style = 'Table Grid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = '                            מק"ט מערכת'
hdr_Cells[1].text = '                             שם ההרכבה'
hdr_Cells[2].text = '                            שם הפרויקט'

for text1, text2, text3 in records:
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text = text1
    row_Cells[1].text = text2
    row_Cells[2].text = text3

menuTable.cell(3, 0).merge(menuTable.cell(4, 0))
menuTable.cell(3, 2).merge(menuTable.cell(4, 2))
menuTable.cell(3, 1).merge(menuTable.cell(4, 1)).merge(menuTable.cell(5, 1)).merge(menuTable.cell(6, 1))
doc.add_paragraph("                                                                                                                                                                                                      בדיקה ויזואליות חיצוניות" )


# Table 2
Table2 = doc.add_table(18, 7)
#Table2.alignment = WD_TABLE_ALIGNMENT.CENTER
Table2.style = 'Table Grid'
Table2.cell(0, 0).text = "                  הערות"
Table2.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table2.cell(0, 0).width = Cm(5.5)
Table2.cell(1, 0).width = Cm(5.5)

Table2.cell(2, 0).text = "   נקה את כלל חלקי מארז     " \
                         "                 המערכת"
Table2.cell(3, 0).text= "     העזר בפנס לאיתור 8 ברגי " \
                        "                   מדחס"

"""a = doc.add_paragraph("ישראלי/ אירופה/ארה''ב ")
a.style = doc.styles.add_style('Style Name', WD_STYLE_TYPE.PARAGRAPH)
font = a.style.font
font.bold = True
font.size = Pt(9)"""
Table2.cell(4, 0).text = "  ישראלי/ אירופה/ארה''ב"
Table2.cell(4, 0).height = Cm(0.5)
Table2.cell(6, 0).text = "         וודא 9 ברגים בגב " \
                         "                          המערכת"
Table2.cell(8, 0).text = "בהתאם לתצור 110/230"
Table2.cell(10, 0).text = "מיקום המסננים בתא המסננים " \
                          "                בגב המכשיר"
Table2.cell(12, 0).text = "       מיקום בידית הנשיאה"
Table2.cell(15, 0).text = "  מדי ספיקה , CB, מתג הפעלה" \
                          "         LCD                "




Table2.cell(0, 1).text = "               אימות תיקון"
Table2.cell(0, 1).width = Cm(3.0)
Table2.cell(0, 1).merge(Table2.cell(0, 2))
Table2.cell(0, 0).merge(Table2.cell(1, 0))
Table2.cell(1, 1).text = "    לא תקין"
Table2.cell(1, 1).width = Cm(1.7)
Table2.cell(1, 2).text = "     תקין"
Table2.cell(1, 2).width = Cm(1.3)
Table2.cell(1, 3).text = "    לא תקין"
Table2.cell(1, 3).width = Cm(1.7)
Table2.cell(1, 4).text = "     תקין"
Table2.cell(1, 4).width = Cm(1.3)
Table2.cell(0, 3).text = "                   בדיקה"
Table2.cell(0, 3).width = Cm(3.0)
Table2.cell(0, 3).merge(Table2.cell(0, 4))
Table2.cell(0, 4).width = Cm(2.9)
Table2.cell(0, 5).text = "                         בדיקה"
Table2.cell(0, 5).width = Cm(7.3)
Table2.cell(0, 5).merge(Table2.cell(1, 5))
Table2.cell(1, 5).width = Cm(7.3)
Table2.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


Table2.cell(2, 6).text = '        1'
Table2.cell(0, 6).width = Cm(1.0)
Table2.cell(1, 6).width = Cm(1.0)
Table2.cell(3, 6).text = '        2'
Table2.cell(4, 6).text = '        3'
Table2.cell(5, 6).text = '        4'
Table2.cell(6, 6).text = '        5'
Table2.cell(7, 6).text = '        6'
Table2.cell(8, 6).text = '        7'
Table2.cell(9, 6).text = '        8'
Table2.cell(10, 6).text = '       9'
Table2.cell(11, 6).text = '      10'
Table2.cell(12, 6).text = '      11'
Table2.cell(13, 6).text = '      12'
Table2.cell(14, 6).text = '      13'
Table2.cell(15, 6).text = '      14'
Table2.cell(16, 6).text = '      15'
Table2.cell(17, 6).text = '      16'
Table2.cell(2, 5).text = "            עטה כפפות הגנה ונקה מערכת " \
                         "                  באמצעות מטלית חיטוי"


Table2.cell(3, 5).text = "     וודא בולמים מחוברים ומהודקים"
Table2.cell(4, 5).text = "            וודא תקע מערכת בהתאם" \
                         "                              לצורת היעד"
Table2.cell(4, 5).font = Pt(5)
Table2.cell(5, 5).text = "    חבר מערכת למקור מתח והפעל על" \
                         "              ספיקה של 5 ליטר לדקה"
Table2.cell(6, 5).text = "        וודא הימצאות כלל ברגיי המארז"
Table2.cell(7, 5).text = "      וודא הימצאות מדבקתייצור כסופה"
Table2.cell(8, 5).text = "                     וודא מדבקת מתחים POWER SUPPLY"
Table2.cell(9, 5).text = "        וודא מדבקת מספר סידורי צבאן " \
                         "                   משולבת ברקוד"
Table2.cell(10, 5).text = "      וודא חיבור מסנן גס, מסנן " \
                          "                       בקטריאלי ומסנן מדחס"
Table2.cell(11, 5).text = "              וודא צנרת מסננים ללא " \
                          "                            כיפופים ושבירות"
Table2.cell(12, 5).text = "           וודא רמקול מערכת מחובר " \
                          "                         וממקם כראוי"
Table2.cell(13, 5).text = "            לוגו צבאן מודפס כראוי"
Table2.cell(14, 5).text = "                     וודא הימצאות מדבקת NO SMOKE"
Table2.cell(15, 5).text = "            וודא כיתוב רכיבים תקין"
Table2.cell(16, 5).text = "  הדבק מדבקת אזהרת ספיקה מירבית"
Table2.cell(17, 5).text = "        הדבק מדבקת החלפת מסננים"

doc.add_paragraph("                                                                                       בדיקת תפעולית - בדיקה עם צב''ד ידני (כחול): במידה ולא בוצע תקין לא נידרש למלא אימות תקין" )

# Table 3

Table3 = doc.add_table(13, 7)
Table3.style = 'Table Grid'



Table3.cell(0, 0).text = "                     הערות"

Table3.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table3.cell(0, 0).width = Cm(6.0)
Table3.cell(0, 1).text = "               אימות תיקון"
Table3.cell(0, 1).merge(Table3.cell(0, 2))
Table3.cell(1, 1).width = Cm(1.7)
Table3.cell(0, 3).text = "                   בדיקה"
Table3.cell(0, 3).merge(Table3.cell(0, 4))
Table3.cell(0, 5).text = "                        בדיקה"
Table3.cell(0, 5).width = Cm(7.5)
Table3.cell(0, 5).merge(Table3.cell(1, 5))
Table3.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table3.cell(0, 6).merge(Table3.cell(1, 6))
Table3.cell(1, 1).text = "     לא תקין"
Table3.cell(1, 2).text = "       תקין"
Table3.cell(1, 3).text = "     לא תקין"
Table3.cell(1, 4).text = "       תקין"

Table3.cell(2, 6).text = "        1"
Table3.cell(2, 5).text = "וודא כי תצוגת המערכת כראוי" \
                         "                                 LCD במסך" \

Table3.cell(3, 6).text = "        2"
Table3.cell(3, 5).text = "         וודא נורות חיווי דולקות"
Table3.cell(4, 6).text = "        3"
Table3.cell(4, 5).text = "כוון מד זרימה לספיקה מערכת של 5 ליטרים." \
                         "חבר מד ידני וודא ספיקה של 5.5 - 4.5 " \
                         "     ליטרים לדקה"
Table3.cell(5, 6).text = "        4"
Table3.cell(5, 5).text = "  וודא קריאת חמצו במד ידני" \
                         " 96% -92%               " \

Table3.cell(6, 6).text = "        5"
Table3.cell(6, 5).text = "     8.5 +- 0.8 סגור מד ידני וודא לחץ"

Table3.cell(7, 6).text = "        6"
Table3.cell(7, 5).text = "סגור מוצא מד ידני וודא התרעות ספיקה " \
                         "OXYGEN FLOW LOW נמוכה "
Table3.cell(8, 6).text = "        7"
Table3.cell(8, 5).text = "סגור מוצא מד ידני וודא כדורית ספיקה " \
                         "                  נופלת לאפס"
Table3.cell(9, 6).text = "        8"
Table3.cell(9, 5).text = "כוון בורג מד ספיקה למקסימום וודא ספיקה " \
                         "    מרבית: 7.5 +- 2 לטרים לדקה"
Table3.cell(10, 6).text = "       9"
Table3.cell(10, 5).text = "             כוון בורג מד ספיקה למקסימום " \
                          "                  וודא התרעת ספיקה גבועה           HIGH OXYGEN FLOW      " \

Table3.cell(11, 6).text = "     10"
Table3.cell(11, 5).text = "עם ירידת החמצן אל מתחת ל 80% הורד " \
                          "ספיקת מערכת 4 ליטרים וודא התרעת " \
                          "חמצן נמוך                                                                " \
                          " OXYGEN LOW"
Table3.cell(12, 6).text = "     11"
Table3.cell(12, 5).text = "נתק מערכת ממקור מתח בעבודה ווודא " \
                          "               התרעת תקלת מתח"
Table3.cell(2, 0).text = "וודא תצוגת מונה שעות מערכת "
Table3.cell(3, 0).text = "         (דלוקות) POWER, NORMAL OXYGEN"\
                         "" \
                         "         (כבויה)     SERVICE REQUIRED     "

Table3.cell(4, 0).text = "מרכז הכדורית מייצג את הספיקה" \
                         "                 הנבחרת"

Table3.cell(7, 0).text = " מהבהבת ושם SERVICE וודא " \
                         "LCD ההתררע מוצג במסך "

Table3.cell(8, 0).text = "           בדיקת נזילות      "
Table3.cell(10, 0).text = " מהבהבת ושם SERVICE וודא " \
                         "LCD ההתררע מוצג במסך "
Table3.cell(11, 0).text = " מהבהבת ושם SERVICE וודא " \
                         "LCD ההתררע מוצג במסך "
Table3.cell(12, 0).text = " מהבהבת ושם SERVICE וודא " \
                         "LCD ההתררע מוצג במסך "

doc.add_paragraph("                                                                                                                                בדיקה פונציונלית - ריצת מערכת עם צב''ד ממוחשב בצורת בדיקה כללית")
# CONTINUAR COM A TABELA 4

Table4 = doc.add_table(10, 9)
Table4.style = 'Table Grid'

Table4.cell(0, 0).text = "               הערות          "
Table4.cell(0, 0).merge(Table4.cell(1, 0))
Table4.cell(0, 0).width = Cm(4.5)
Table4.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table4.cell(4, 0).text = "     חבר צינור צב''ד והפעל  בהתאם לתכנית "

Table4.cell(8, 0).text = "       בדיקת נזילות"
Table4.cell(9, 0).text = "נדרש לוודא כי האוויר" \
                         "            אינו חם"


Table4.cell(0, 1).text = "          אימות תיקון"
Table4.cell(0, 1).merge(Table4.cell(0, 2))

Table4.cell(1, 1).text = "     לא תקין"
Table4.cell(1, 2).text = "     תקין"
Table4.cell(1, 4).text = "     לא תקין"
Table4.cell(1, 5).text = "     תקין"

Table4.cell(0, 3).text = "      ערך"
Table4.cell(0, 4).text = "               בדיקה"
Table4.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table4.cell(0, 4).merge(Table4.cell(0, 5))
Table4.cell(0, 6).text = "     ערך"
Table4.cell(0, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table4.cell(0, 6).merge(Table4.cell(1, 6))
Table4.cell(0, 7).text = "                    בדיקה"
Table4.cell(0, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table4.cell(0, 7).merge(Table4.cell(1, 7))
Table4.cell(0, 7).width = Cm(6.3)

Table4.cell(2, 7).text = " LPM ספיקת מערכת בצב''ד "
Table4.cell(3, 7).text = "       % רמת חמצן במוצא"
Table4.cell(4, 7).text = "חבר מערכת למקור מתח, הפעל וכוון ספיקה ל5 ליטרים )4.5 -5.5)"
Table4.cell(5, 7).text = "        LPM תנודתיות ספיקה"
Table4.cell(6, 7).text = "       תנודתיות אחוז חמצן בגרף"
Table4.cell(7, 7).text = "               PSI לחץ מוצר"
Table4.cell(8, 7).text = "סגור מוצא מד ידני וודא כדורית              ספיקה נופלת לאפס"
Table4.cell(9, 7).text = "וודא טמפרטורת מוצא מחולל                            קרירה"

Table4.cell(4, 3).text = "    N/A"
Table4.cell(4, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

Table4.cell(8, 3).text = "    N/A"
Table4.cell(8, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

Table4.cell(9, 3).text = "    N/A"
Table4.cell(9, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

Table4.cell(4, 6).text = "    N/A"
Table4.cell(4, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

Table4.cell(8, 6).text = "    N/A"
Table4.cell(8, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

Table4.cell(9, 6).text = "    N/A"
Table4.cell(9, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

Table4.cell(2, 8).text = "      1"
Table4.cell(3, 8).text = "      2"
Table4.cell(4, 8).text = "      3"
Table4.cell(5, 8).text = "      4"
Table4.cell(6, 8).text = "      5"
Table4.cell(7, 8).text = "      6"
Table4.cell(8, 8).text = "      7"
Table4.cell(9, 8).text = "      8"

doc.add_paragraph("                                                                                                                                                                                                                     בדיקות אריזה")
# Table 5
Table5 = doc.add_table(11, 7)
Table5.style = 'Table Grid'

Table5.cell(0, 0).text = "                    הערות"
Table5.cell(0, 0).width = Cm(6.0)

Table5.cell(0, 1).text = "               אימות תקין"
Table5.cell(0, 1).merge(Table5.cell(0, 2))

Table5.cell(0, 3).text = "                 בדיקה"
Table5.cell(0, 3).width = Cm(1.95)

Table5.cell(0, 4).width = Cm(1.05)
Table5.cell(0, 3).merge(Table5.cell(0, 4))
Table5.cell(0, 5).text = "                            בדיקה"
Table5.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
Table5.cell(0, 5).merge(Table5.cell(1, 5))
Table5.cell(0, 5).width = Cm(8.0)
Table5.cell(1, 1).text = "   לא תקין"
Table5.cell(1, 1).width = Cm(1.95)
Table5.cell(1, 2).text = "     תקין"
Table5.cell(1, 2).width = Cm(1.05)
Table5.cell(1, 3).text = "   לא תקין"
Table5.cell(1, 3).width = Cm(1.95)
Table5.cell(1, 4).text = "     תקין"
Table5.cell(1, 4).width = Cm(1.05)

Table5.cell(2, 0).text = "        באם מלוכלכת יש " \
                         "                       לנקות"

Table5.cell(2, 5).text = "וודא מערכת ללא פגיעות, שריטות ולכלוך"
Table5.cell(3, 0).text = "       באתאם לתצורה"
Table5.cell(4, 0).text = "       באתאם לתצורה"
Table5.cell(5, 0).text = "       באתאם לתצורה"
Table5.cell(8, 0).text = "     וודא כי מתאים למספרו הסידורי של " \
                         " המכשיר"
Table5.cell(9, 0).text = "     חייב כי מתאים למספרו הסידורי של " \
                         "  המכשיר"
Table5.cell(10, 0).text = " ביצוע ע''י אורז מערכות לאימות תקינות" \
                          " אריזה ותקינות" \
                          "   מילוי הטופס"

Table5.cell(3, 5).text = "        סוג אריזה תואם להזמנה"
Table5.cell(4, 5).text = "על גבי האריזה SHIPMENT וודא מדבקה "
Table5.cell(5, 5).text = "וודא הוראות שימוש תואמות מצורפות"
Table5.cell(6, 5).text = "וודא שתי צינורות אף מצורפות"
Table5.cell(7, 5).text = "    (HUMIDIFIER)וודא שילוב מיכל " \
                         "                                  אידוי מצורף "
Table5.cell(8, 5).text = "  צבאן ע''ג האריזה SN הדבק הדבקה"
Table5.cell(9, 5).text = "הכנס מספר סידורי האריזה בשדה הברקוד " \
                         "                                        בכותרת"
Table5.cell(10, 5).text = "    בדיקת התאמת מדבקות סידורי ע''י " \
                          "                                    בודק שני"
Table5.cell(2, 6).text = "       1"
Table5.cell(3, 6).text = "       2"
Table5.cell(4, 6).text = "       3"
Table5.cell(5, 6).text = "       4"
Table5.cell(6, 6).text = "       5"
Table5.cell(7, 6).text = "       6"
Table5.cell(8, 6).text = "       7"
Table5.cell(9, 6).text = "       8"
Table5.cell(10, 6).text = "      9"
doc.add_paragraph("")
doc.add_paragraph("________________________תאריך                                                                                          שם בודק                                                             חתימה             ")
doc.add_paragraph("________________________תאריך                                                                                      שם בודק שני                                                           חתימה             ")

doc.save('table.docx')
os.system("start table.docx")